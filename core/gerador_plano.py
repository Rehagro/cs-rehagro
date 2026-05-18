import os
import sys
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from core.mapeamento import MODULO_BOASVINDAS, get_dor_por_id

_LOGO_VERDE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "assets", "logo_rehagro_verde.png"
)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    tcPr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str,
                   color_hex: str = "C9A84C", size_pt: int = 11):
    """Insere um hyperlink clicável no parágrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex.replace("#", ""))
    rPr.append(color_el)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(size_pt * 2))  # half-points
    rPr.append(size_el)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_heading(doc: Document, text: str, level: int = 1,
                 color_hex: str = "1C3829", size_pt: int = 16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return p


def _add_paragraph(doc: Document, text: str, size_pt: int = 11,
                   bold: bool = False, color_hex: str = "1A1A1A",
                   space_before: int = 0, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return p


def _add_modulo_block(doc: Document, ordem: str, modulo: dict,
                      descricao_intro: str = ""):
    """Adiciona um bloco de módulo ao documento."""
    # Linha separadora
    doc.add_paragraph("─" * 60).paragraph_format.space_after = Pt(4)

    # Ordem e nome do módulo
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f"{ordem}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)  # gold

    r2 = p.add_run(modulo["modulo"])
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x1C, 0x38, 0x29)  # green

    # Texto introdutório
    if descricao_intro:
        _add_paragraph(doc, descricao_intro, size_pt=11, space_before=2, space_after=4)

    # Link
    if modulo.get("link"):
        p_link = doc.add_paragraph()
        p_link.paragraph_format.space_before = Pt(2)
        p_link.paragraph_format.space_after = Pt(2)
        r_label = p_link.add_run("Link de acesso: ")
        r_label.bold = True
        r_label.font.size = Pt(11)

        link_val = modulo["link"]
        if link_val.startswith("http://") or link_val.startswith("https://"):
            # Hyperlink clicável (azul ou dourado, sublinhado)
            _add_hyperlink(p_link, link_val, link_val, color_hex="0563C1", size_pt=11)
        else:
            # Texto comum (placeholder/nome do curso)
            r_link = p_link.add_run(link_val)
            r_link.font.size = Pt(11)
            r_link.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

    # Detalhes
    details = []
    if modulo.get("aulas"):
        details.append(f"Quantidade de aulas: {modulo['aulas']}")
    if modulo.get("tempo"):
        details.append(f"Tempo de aula gravada aproximado: {modulo['tempo']}")
    if modulo.get("atividades"):
        details.append(f"Atividades para realizar: {modulo['atividades']}")
    if modulo.get("programacao"):
        details.append(
            f"Programação: Se programe para assistir esse conteúdo em {modulo['programacao']}"
        )

    if details:
        p_mat = doc.add_paragraph()
        p_mat.paragraph_format.space_before = Pt(4)
        p_mat.paragraph_format.space_after = Pt(2)
        r_mat = p_mat.add_run("Materiais de aula:")
        r_mat.bold = True
        r_mat.font.size = Pt(11)

    for detail in details:
        p_d = doc.add_paragraph(style="List Bullet")
        p_d.paragraph_format.space_before = Pt(1)
        p_d.paragraph_format.space_after = Pt(1)
        run = p_d.add_run(detail)
        run.font.size = Pt(10)

    doc.add_paragraph()  # espaço


def gerar_plano_docx(resposta: dict) -> bytes:
    """
    Gera o plano de aula personalizado em .docx e retorna como bytes.
    """
    nome = resposta.get("nome", "Aluno")
    p1_id = resposta.get("prioridade_1")
    p2_id = resposta.get("prioridade_2")
    p3_id = resposta.get("prioridade_3")

    mod1 = get_dor_por_id(p1_id) if p1_id else None
    mod2 = get_dor_por_id(p2_id) if p2_id else None
    mod3 = get_dor_por_id(p3_id) if p3_id else None

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Logo Rehagro (verde) ──────────────────────────────
    if os.path.exists(_LOGO_VERDE_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(8)
        run_logo = p_logo.add_run()
        run_logo.add_picture(_LOGO_VERDE_PATH, width=Inches(1.8))

    # ── Título ───────────────────────────────────────────
    _add_heading(doc, f"Plano de aula — {nome}",
                 level=1, color_hex="1C3829", size_pt=18)

    _add_paragraph(
        doc,
        f"Rehagro · Customer Success  |  {resposta.get('turma_nome', '')}  |  "
        f"Gerado em {datetime.now().strftime('%d/%m/%Y')}",
        size_pt=9, color_hex="6B6B5E", space_after=12
    )

    # ── Introdução ────────────────────────────────────────
    intro = (
        f"{nome}, esse é o seu plano de aula, personalizado de acordo com os seus "
        f"principais desafios. Nele você verá as orientações sobre a sequência de "
        f"conteúdos para assistir nesses primeiros meses de curso, além de materiais "
        f"complementares que podem te apoiar no seu dia a dia. Confira abaixo nossa "
        f"recomendação:"
    )
    _add_paragraph(doc, intro, size_pt=11, space_before=4, space_after=14)

    # ── Módulo: Boas-vindas ───────────────────────────────
    _add_heading(doc, "Para começar", level=2, color_hex="C9A84C", size_pt=13)
    _add_modulo_block(
        doc, "Início",
        MODULO_BOASVINDAS,
        "Para iniciar, veja como funciona o curso e os critérios de aprovação."
    )

    # ── Módulos das 3 dores ───────────────────────────────
    _add_heading(doc, "Sua trilha personalizada", level=2,
                 color_hex="C9A84C", size_pt=13)

    ordens = ["1ª prioridade", "2ª prioridade", "3ª prioridade"]
    intros = [
        "Em seguida, veja o módulo que vai te ajudar com o desafio que você identificou como mais urgente.",
        "Depois de finalizado o módulo anterior, confira este conteúdo.",
        "Por fim, finalize sua trilha com este módulo.",
    ]
    for i, modulo in enumerate([mod1, mod2, mod3]):
        if modulo:
            _add_modulo_block(doc, ordens[i], modulo, intros[i])

    # ── Encerramento ─────────────────────────────────────
    doc.add_paragraph("─" * 60)
    _add_paragraph(
        doc,
        "Nos acione a qualquer momento que precisar! "
        "Estamos aqui para garantir que você tire o máximo deste curso.",
        size_pt=11, space_before=8, space_after=4
    )
    _add_paragraph(doc, "Equipe Customer Success — Rehagro",
                   size_pt=11, bold=True, color_hex="1C3829")

    # Salvar em bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
